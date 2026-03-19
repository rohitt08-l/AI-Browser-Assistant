from fastapi import FastAPI
from pydantic import BaseModel
from groq import Groq
from langchain_text_splitters import RecursiveCharacterTextSplitter 
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from docx import Document
import os
from dotenv import load_dotenv

load_dotenv()

api_key = os.getenv("GROQ_API_KEY")
client = Groq(api_key=api_key)

app = FastAPI()

embedding_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

class RequestData(BaseModel):
    question: str
    page_content: str
    task_type: str = "auto"   # new field


# ---------------- LLM ----------------
def groq_llm(prompt):
    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content


# ---------------- TASK CLASSIFIER ----------------
def classify_task(question):
    prompt = f"""
    Classify the user's intent into:
    QA, SUMMARIZE, NOTES, GENERATE_DOC

    Question: {question}

    Only return one word.
    """
    return groq_llm(prompt).strip()


# ---------------- TASK HANDLERS ----------------
def handle_qa(context, question):
    return groq_llm(f"""
    Answer clearly using context only.

    Context:
    {context}

    Question:
    {question}
    """)


def handle_summary(context):
    return groq_llm(f"""
    Summarize the following in simple points:

    {context}
    """)


def handle_notes(context):
    return groq_llm(f"""
    Create structured notes with headings, bullet points and key insights:

    {context}
    """)


def handle_doc(context):
    return groq_llm(f"""
    Create a professional document format with title, headings and paragraphs:

    {context}
    """)


# ---------------- DOCX GENERATOR ----------------
def create_docx(content):
    doc = Document()
    doc.add_heading('Generated Document', 0)
    doc.add_paragraph(content)

    file_path = "output.docx"
    doc.save(file_path)

    return file_path


# ---------------- MAIN API ----------------
@app.post("/ask")
def ask(data: RequestData):

    # Split content
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100
    )
    chunks = splitter.split_text(data.page_content[:10000])

    # Vector DB
    vectorstore = FAISS.from_texts(chunks, embedding_model)

    docs = vectorstore.similarity_search(data.question, k=3)
    context = "\n\n".join([doc.page_content for doc in docs])

    # Task selection
    task = data.task_type.upper()

    if task == "AUTO":
        task = classify_task(data.question)

    # Route tasks
    if task == "QA":
        answer = handle_qa(context, data.question)
        return {"answer": answer}

    elif task == "SUMMARIZE":
        answer = handle_summary(context)
        return {"answer": answer}

    elif task == "NOTES":
        answer = handle_notes(context)
        return {"answer": answer}

    elif task == "GENERATE_DOC":
        content = handle_doc(context)
        file_path = create_docx(content)
        return {"answer": content, "file": file_path}

    else:
        return {"answer": "Task not recognized"}